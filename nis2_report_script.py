import os
import re
import shutil
import subprocess
from datetime import datetime
from docx import Document
from docx.oxml import parse_xml

# === CONFIGURATION ===

TEMPLATES = {
    "fr": {
        "important": "templates/IMPORTANT_NIS2_TEMPLATE_FR.docx",
        "basic": "templates/BASIC_NIS2_TEMPLATE_FR.docx",
        "essentiel": "templates/ESSENTIAL_NIS2_TEMPLATE_FR.docx"
    },
    "en": {
        "important": "templates/IMPORTANT_NIS2_TEMPLATE_EN.docx",
        "basic": "templates/BASIC_NIS2_TEMPLATE_EN.docx",
        "essentiel": "templates/ESSENTIAL_NIS2_TEMPLATE_EN.docx"
    },
    "nl": {
        "important": "templates/IMPORTANT_NIS2_TEMPLATE_NL.docx",
        "basic": "templates/BASIC_NIS2_TEMPLATE_NL.docx",
        "essentiel": "templates/ESSENTIAL_NIS2_TEMPLATE_NL.docx"
    }
}

STATUT_COLORS = {
    "Conforme": "C6EFCE", "Non conforme": "FFC7CE", "Partiellement conforme": "FFEB9C",
    "Non applicable": "D9D9D9", "Non évalué": "D9D9D9",
    "Compliant": "C6EFCE", "Non-compliant": "FFC7CE", "Partially compliant": "FFEB9C",
    "Not applicable": "D9D9D9", "Not assessed": "D9D9D9",
    "Conform": "C6EFCE", "Niet conform": "FFC7CE", "Gedeeltelijk conform": "FFEB9C",
    "Niet van toepassing": "D9D9D9", "Niet beoordeeld": "D9D9D9"
}

STATUT_MAPPING_FR = {
    "Compliant": "Conforme", "Non-compliant": "Non conforme", "Partially compliant": "Partiellement conforme",
    "Not applicable": "Non applicable", "Not assessed": "Non évalué",
    "Conforme": "Conforme", "Non conforme": "Non conforme", "Partiellement conforme": "Partiellement conforme",
    "Non applicable": "Non applicable", "Non évalué": "Non évalué"
}

STATUT_MAPPING_EN = {
    "Compliant": "Compliant", "Non-compliant": "Non-compliant", "Partially compliant": "Partially compliant",
    "Not applicable": "Not applicable", "Not assessed": "Not assessed",
    "Conforme": "Compliant", "Non conforme": "Non-compliant", "Partiellement conforme": "Partially compliant",
    "Non applicable": "Not applicable", "Non évalué": "Not assessed"
}

STATUT_MAPPING_NL = {
    "Compliant": "Conform", "Non-compliant": "Niet conform", "Partially compliant": "Gedeeltelijk conform",
    "Not applicable": "Niet van toepassing", "Not assessed": "Niet beoordeeld",
    "Conforme": "Conform", "Non conforme": "Niet conform", "Partiellement conforme": "Gedeeltelijk conform",
    "Non applicable": "Niet van toepassing", "Non évalué": "Niet beoordeeld"
}

# === UTILS ===

def format_balise(text):
    return re.sub(r'[^a-zA-Z0-9]', '_', text).upper()

def extract_group(text):
    match = re.search(r"(basic|important|essentiel|essential)", text, re.IGNORECASE)
    if match:
        grp = match.group(1).lower()
        return "essentiel" if grp == "essential" else grp
    return "basic"

def extract_title(text):
    match = re.search(r"Audit[_\s]?Nis2[_\s]?([A-Z0-9_-]+)", text, re.IGNORECASE)
    return match.group(0).replace(" ", "_") if match else ""

def clean_title(title):
    return re.sub(r"Audit[_\s]?Nis2[_\s]?", "", title, flags=re.IGNORECASE).strip()

def extract_compliance(text, lang):
    lines = text.split("\n")
    results = []
    current = {"id": None, "status": "", "observation": ""}
    in_obs = False

    statut_map = STATUT_MAPPING_EN if lang == "en" else STATUT_MAPPING_NL if lang == "nl" else STATUT_MAPPING_FR
    is_status = lambda t: t in statut_map
    format_status = lambda t: statut_map.get(t, t)

    for line in lines:
        trimmed = line.strip()
        if re.match(r"^(BASIC_|IMPORTANT_)?[A-Z]{2}\.[A-Z]{2,}-\d+(\.\d+)?$", trimmed) or trimmed in ["R.AC-3.5", "R.AC-3.4"]:
            if current["id"]:
                results.append(current.copy())
            current = {"id": trimmed, "status": "", "observation": ""}
            in_obs = False
        elif is_status(trimmed):
            current["status"] = format_status(trimmed)
            in_obs = False
        elif re.match(r"^Observation\s*:", trimmed, re.IGNORECASE):
            current["observation"] = re.sub(r"^Observation\s*:", "", trimmed, flags=re.IGNORECASE).strip()
            in_obs = True
        elif in_obs and trimmed:
            current["observation"] += " " + trimmed
        elif in_obs and not trimmed:
            in_obs = False

    if current["id"]:
        results.append(current)

    return results

def set_cell_color(cell, hex_color):
    shading_elm = parse_xml(f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def replace_placeholders(doc, replacements, color_targets):
    found_tags = set()
    missing_tags = set(replacements.keys())

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, val in replacements.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, val)
                        found_tags.add(key)
                        if key in color_targets:
                            set_cell_color(cell, color_targets[key])

    for p in doc.paragraphs:
        for key, val in replacements.items():
            if key in p.text:
                p.text = p.text.replace(key, val)
                found_tags.add(key)

    missing_tags -= found_tags
    print(f"🟢 {len(found_tags)} balises remplacées avec succès.")
    if missing_tags:
        print(f"🟡 {len(missing_tags)} balises non trouvées :")
        for tag in missing_tags:
            print(f"   ⛔ {tag}")

def convert_docx_to_pdf(input_path, output_dir):
    try:
        subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf", input_path, "--outdir", output_dir],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE
        )
        pdf_name = os.path.splitext(os.path.basename(input_path))[0] + ".pdf"
        print(f"📄 PDF généré : {os.path.join(output_dir, pdf_name)}")
    except subprocess.CalledProcessError as e:
        print("❌ Erreur lors de la conversion en PDF :", e.stderr.decode())

def generate_report(source_path, output_dir, custom_title=None, lang="fr", client_name=""):
    print("📥 Lecture du fichier source...")
    source_doc = Document(source_path)
    full_text = "\n".join([p.text for p in source_doc.paragraphs])
    group = extract_group(full_text)
    template_path = TEMPLATES[lang][group]
    print(f"🔧 Modèle utilisé : {group} ({lang})")

    title_raw = extract_title(full_text)
    title_clean = clean_title(title_raw)

    if not custom_title:
        custom_title = f"Rapport_NIS2_{title_clean or 'Généré'}"

    today = datetime.now().strftime("%d/%m/%Y")
    os.makedirs(output_dir, exist_ok=True)
    output_docx = os.path.join(output_dir, f"{custom_title}.docx")

    print(f"📄 Génération du fichier Word : {output_docx}")
    shutil.copy(template_path, output_docx)
    output_doc = Document(output_docx)

    compliance_data = extract_compliance(full_text, lang)
    replacements = {
        "{{DATE}}": today,
        "{{TITLE}}": title_clean,
        "{{CLIENT}}": client_name  # <- ajout ici
    }
    color_targets = {}

    for item in compliance_data:
        key = format_balise(item["id"])
        statut = item["status"]
        replacements[f"{{{{STATUT_{key}}}}}"] = statut
        replacements[f"{{{{OBSERVATION_{key}}}}}"] = item["observation"]
        if statut in STATUT_COLORS:
            color_targets[f"{{{{STATUT_{key}}}}}"] = STATUT_COLORS[statut]

    replace_placeholders(output_doc, replacements, color_targets)
    output_doc.save(output_docx)
    print("✅ Fichier Word généré avec succès :", output_docx)
    convert_docx_to_pdf(output_docx, output_dir)
    print("✅ Rapport complet disponible dans :", output_dir)

# === MAIN ===

if __name__ == "__main__":
    try:
        audit_dir = "./audits"
        files = [f for f in os.listdir(audit_dir) if f.endswith(".html") or f.endswith(".docx")]

        if not files:
            print("❌ Aucun fichier .html ou .docx trouvé dans ./audits")
            exit(1)

        print("📄 Sélectionnez un fichier à traiter :")
        for i, file in enumerate(files, start=1):
            print(f"  [{i}] {file}")

        choix = input("🔢 Entrez le numéro du fichier à traiter : ").strip()
        if not choix.isdigit() or not (1 <= int(choix) <= len(files)):
            print("❌ Sélection invalide.")
            exit(1)

        selected = files[int(choix)-1]
        selected_path = os.path.join(audit_dir, selected)

        if selected.endswith(".html"):
            docx_name = os.path.splitext(selected)[0] + ".docx"
            docx_path = os.path.join(audit_dir, docx_name)
            print(f"🔄 Conversion HTML vers DOCX via LibreOffice : {selected} → {docx_name}")
            try:
                subprocess.run([
                    "libreoffice", "--headless", "--convert-to", "odt", "--outdir", audit_dir, selected_path
                ], check=True)
                odt_path = os.path.join(audit_dir, os.path.splitext(selected)[0] + ".odt")
                subprocess.run([
                    "libreoffice", "--headless", "--convert-to", "docx", "--outdir", audit_dir, odt_path
                ], check=True)
                os.remove(selected_path)
                os.remove(odt_path)
                print(f"✅ Fichier converti avec succès : {docx_path}")
                selected_path = docx_path
            except subprocess.CalledProcessError as e:
                print("❌ Erreur LibreOffice :", e)
                exit(1)

        lang = input("🌐 Langue du template (fr/en/nl) [fr] : ").strip().lower() or "fr"
        if lang not in ["fr", "en", "nl"]:
            print("❌ Langue invalide. Utilisez 'fr', 'en' ou 'nl'.")
            exit(1)

        custom_title = input("📝 Nom du fichier de sortie (laisser vide pour générer automatiquement) : ").strip()
        client_name = input("👤 Nom du client (sera utilisé pour remplacer {{CLIENT}}) : ").strip()
        output_dir = f"./output/{lang}"

        generate_report(selected_path, output_dir, custom_title or None, lang, client_name)

    except Exception as e:
        print(f"❌ Erreur critique : {str(e)}")
